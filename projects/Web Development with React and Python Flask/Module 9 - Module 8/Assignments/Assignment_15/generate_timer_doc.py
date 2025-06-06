
from docx import Document

def create_timer_documentation(filename='Timer_Component_Documentation.docx'):
    doc = Document()
    doc.add_heading('Timer Component Documentation', 0)

    sections = {
        "Overview": "This timer component manages countdown functionality for each test section using the `react-timer-hook` library. "
                    "It supports start, pause, and reset operations, and is designed to be reusable across test sections.",

        "Installation": "Install react-timer-hook:\n\n"
                        "```bash\nnpm install react-timer-hook\n```",

        "Timer Component (TestTimer.jsx)": "```jsx\n"
            "import React from \"react\";\n"
            "import { useTimer } from \"react-timer-hook\";\n\n"
            "function TestTimer({ durationSeconds, onExpire }) {\n"
            "  const expiryTimestamp = new Date();\n"
            "  expiryTimestamp.setSeconds(expiryTimestamp.getSeconds() + durationSeconds);\n\n"
            "  const {\n"
            "    seconds,\n"
            "    minutes,\n"
            "    isRunning,\n"
            "    start,\n"
            "    pause,\n"
            "    resume,\n"
            "    restart,\n"
            "  } = useTimer({\n"
            "    expiryTimestamp,\n"
            "    onExpire: () => {\n"
            "      console.warn(\"Timer expired\");\n"
            "      if (onExpire) onExpire();\n"
            "    },\n"
            "  });\n\n"
            "  return (\n"
            "    <div className=\"timer-container\">\n"
            "      <p>Time Remaining: {minutes}:{seconds}</p>\n"
            "      <button onClick={pause}>Pause</button>\n"
            "      <button onClick={resume}>Resume</button>\n"
            "      <button onClick={() => restart(expiryTimestamp)}>Reset</button>\n"
            "    </div>\n"
            "  );\n"
            "}\n\n"
            "export default TestTimer;\n"
            "```",

        "Features": "- Countdown from a set duration\n"
                    "- Visual updates when time is running low (add styles as needed)\n"
                    "- Pause, resume, and reset functionality\n"
                    "- Callback on timer expiry",

        "Integration Tips": "- Place the `TestTimer` component near your question UI\n"
                            "- Sync timer resets with question changes\n"
                            "- Use state to manage active sections",

        "Accessibility Considerations": "- Use ARIA labels for screen readers\n"
                                       "- Provide color cues and text warnings for low time"
    }

    for title, content in sections.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    doc.save(filename)
    print(f"Documentation saved as {filename}")

if __name__ == "__main__":
    create_timer_documentation()
